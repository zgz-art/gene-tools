import streamlit as st

# 必须是第一个 Streamlit 命令
st.set_page_config(page_title="简历智能填充工具", page_icon="📄", layout="wide")

# 然后导入其他所有库
import os
import sys
import base64
import io
import tempfile
import json
import re
from datetime import datetime
from copy import copy

# 强制 stdout/stderr 使用 UTF-8
sys.stdout.reconfigure(encoding='utf-8')
sys.stderr.reconfigure(encoding='utf-8')
os.environ["PYTHONIOENCODING"] = "utf-8"
os.environ["LANG"] = "zh_CN.UTF-8"

import pypdf
from zhipuai import ZhipuAI
import openpyxl
from docx import Document
from docx.shared import Inches, Emu
from PIL import Image
import numpy as np

# 尝试导入 RapidOCR
try:
    from rapidocr_onnxruntime import RapidOCR
    RAPIDOCR_AVAILABLE = True
except ImportError:
    RAPIDOCR_AVAILABLE = False
    # 注意：此时可以安全使用 st.warning，因为 set_page_config 已经执行

if not RAPIDOCR_AVAILABLE:
    st.warning("⚠️ RapidOCR 未安装，Word 填充功能将不可用。请运行：pip install rapidocr-onnxruntime")

# ==================== 页面样式 ====================
st.markdown("""
<style>
    .stButton button { background-color: #4CAF50; color: white; border-radius: 8px; }
    h1 { color: #2c3e50; text-align: center; }
    .score-card {
        background-color: #f0f2f6;
        border-radius: 12px;
        padding: 20px;
        margin: 10px 0;
    }
    .risk-card {
        background-color: #fff3e0;
        border-left: 5px solid #ff9800;
        border-radius: 8px;
        padding: 15px;
        margin: 10px 0;
    }
</style>
""", unsafe_allow_html=True)

# ==================== 初始化 Session State ====================
if "pdf_content" not in st.session_state:
    st.session_state.pdf_content = None
if "excel_content" not in st.session_state:
    st.session_state.excel_content = None
if "ai_result" not in st.session_state:
    st.session_state.ai_result = None
if "pdf_text" not in st.session_state:
    st.session_state.pdf_text = None
if "last_pdf_name" not in st.session_state:
    st.session_state.last_pdf_name = None
if "last_excel_name" not in st.session_state:
    st.session_state.last_excel_name = None
if "score_result" not in st.session_state:
    st.session_state.score_result = None
if "risk_result" not in st.session_state:
    st.session_state.risk_result = None
if "tech_requirements" not in st.session_state:
    st.session_state.tech_requirements = ""
if "level" not in st.session_state:
    st.session_state.level = ""

# 图片相关状态
if "uploaded_images" not in st.session_state:
    st.session_state.uploaded_images = None
if "image_text" not in st.session_state:
    st.session_state.image_text = None
if "last_images_hash" not in st.session_state:
    st.session_state.last_images_hash = None

# Word 模板相关状态
if "word_template" not in st.session_state:
    st.session_state.word_template = None
if "last_word_name" not in st.session_state:
    st.session_state.last_word_name = None

# ==================== 辅助函数 ====================
def normalize_date(date_str: str) -> str:
    if not date_str:
        return ""
    date_str = str(date_str).strip()
    for fmt in ["%Y-%m-%d", "%Y/%m/%d", "%Y.%m.%d", "%Y年%m月%d日", "%Y-%m", "%Y/%m", "%Y.%m", "%Y年%m月"]:
        try:
            if fmt.endswith("%m"):
                return datetime.strptime(date_str, fmt).strftime("%Y-%m-01")
            else:
                return datetime.strptime(date_str, fmt).strftime("%Y-%m-%d")
        except ValueError:
            continue
    year_match = re.search(r"\b(19|20)\d{2}\b", date_str)
    return f"{year_match.group()}-01-01" if year_match else date_str

def extract_text_from_pdf(pdf_file) -> str:
    reader = pypdf.PdfReader(pdf_file)
    return "\n".join(page.extract_text() for page in reader.pages)

def extract_text_from_images(api_key: str, model: str, image_files) -> str:
    if not image_files:
        return ""
    client = ZhipuAI(api_key=api_key)
    vision_model = "glm-4v-plus"
    all_texts = []
    for img_file in image_files:
        img_bytes = img_file.read()
        img_base64 = base64.b64encode(img_bytes).decode('utf-8')
        mime_type = img_file.type if img_file.type else "image/jpeg"
        data_url = f"data:{mime_type};base64,{img_base64}"
        messages = [
            {
                "role": "user",
                "content": [
                    {"type": "text", "text": "请提取这张图片中的所有文字内容，只返回文字，不要额外解释。如果图片没有文字，返回空字符串。"},
                    {"type": "image_url", "image_url": {"url": data_url}}
                ]
            }
        ]
        try:
            response = client.chat.completions.create(
                model=vision_model,
                messages=messages,
                temperature=0.0,
            )
            text = response.choices[0].message.content.strip()
            if text:
                all_texts.append(f"【来自图片：{img_file.name}】\n{text}")
        except Exception as e:
            st.warning(f"图片 {img_file.name} 识别失败: {e}")
            continue
        img_file.seek(0)
    return "\n\n".join(all_texts)

def get_primary_cell(worksheet, row, col):
    for merged in worksheet.merged_cells.ranges:
        if (merged.min_row <= row <= merged.max_row and
            merged.min_col <= col <= merged.max_col):
            return worksheet.cell(merged.min_row, merged.min_col)
    return worksheet.cell(row, col)

def safe_write(worksheet, row, col, value):
    get_primary_cell(worksheet, row, col).value = value

def normalize_string(s: str) -> str:
    if not s:
        return ""
    s = str(s).strip()
    s = re.sub(r'[：:，,。、；;]', '', s)
    s = re.sub(r'\s+', '', s)
    return s

def find_cell(worksheet, value, exact=False):
    norm_value = normalize_string(value)
    for row in worksheet.iter_rows():
        for cell in row:
            if cell.value is None:
                continue
            if exact:
                if normalize_string(cell.value) == norm_value:
                    return cell.row, cell.column
            else:
                if norm_value in normalize_string(cell.value):
                    return cell.row, cell.column
    return None

def find_row_by_keyword(worksheet, keyword):
    norm_keyword = normalize_string(keyword)
    for row in range(1, worksheet.max_row + 1):
        for col in range(1, worksheet.max_column + 1):
            val = worksheet.cell(row, col).value
            if val and norm_keyword in normalize_string(val):
                return row
    return None

def copy_row_style(source_row, target_row, worksheet, max_col):
    for col in range(1, max_col + 1):
        src = worksheet.cell(source_row, col)
        tgt = worksheet.cell(target_row, col)
        if src.has_style:
            tgt.font = copy(src.font)
            tgt.border = copy(src.border)
            tgt.fill = copy(src.fill)
            tgt.number_format = src.number_format
            tgt.alignment = copy(src.alignment)

def clear_row_content(worksheet, row, max_col):
    for col in range(1, max_col + 1):
        primary = get_primary_cell(worksheet, row, col)
        if primary.row == row and primary.column == col:
            primary.value = None

def fill_basic_info(ws, basic_data):
    for key, value in basic_data.items():
        if not value:
            continue
        pos = find_cell(ws, key, exact=False)
        if pos:
            safe_write(ws, pos[0], pos[1] + 1, value)

def fill_education_block(ws, keyword, edu_data, is_undergrad=True):
    if is_undergrad:
        if "毕业院校" in edu_data and edu_data["毕业院校"]:
            safe_write(ws, 14, 4, edu_data["毕业院校"])
        if "专业" in edu_data and edu_data["专业"]:
            safe_write(ws, 15, 4, edu_data["专业"])
    else:
        if "毕业院校" in edu_data and edu_data["毕业院校"]:
            safe_write(ws, 23, 4, edu_data["毕业院校"])
        if "专业" in edu_data and edu_data["专业"]:
            safe_write(ws, 24, 4, edu_data["专业"])

    pos = find_cell(ws, keyword, exact=True)
    if not pos:
        st.warning(f"未找到关键字: {keyword}")
        return
    row_start = pos[0] + 1
    for offset in range(9):
        current_row = row_start + offset
        field_cell = ws.cell(current_row, 1)
        if field_cell.value is None:
            if offset > 10:
                break
            continue
        field_name = normalize_string(field_cell.value)
        for data_key, data_val in edu_data.items():
            if data_key in ["毕业院校", "专业"]:
                continue
            if data_val and normalize_string(data_key) == field_name:
                safe_write(ws, current_row, 2, data_val)
                break

def fill_work_experience(ws, work_list):
    keyword_row = find_row_by_keyword(ws, "工作经历（由近及远，仅限IT相关经历）")
    if not keyword_row:
        return
    header_row = keyword_row + 1
    data_start = header_row + 1
    reserved = 3
    max_col = ws.max_column

    for r in range(data_start, data_start + reserved):
        clear_row_content(ws, r, max_col)

    sorted_work = sorted(work_list, key=lambda x: x.get("开始日期", "1900-01-01"), reverse=True)
    need = len(sorted_work)
    if need > reserved:
        st.warning(f"工作经历共有 {need} 条，但模板只预留了 {reserved} 行，超出部分将被忽略。请手动增加模板预留行数。")
        need = reserved

    for idx in range(need):
        work = sorted_work[idx]
        target_row = data_start + idx
        for col in range(1, max_col + 1):
            header = ws.cell(header_row, col).value
            if not header:
                continue
            header_norm = normalize_string(header)
            if "开始日期" in header_norm:
                safe_write(ws, target_row, col, normalize_date(work.get("开始日期", "")))
            elif "结束日期" in header_norm:
                safe_write(ws, target_row, col, normalize_date(work.get("结束日期", "")))
            elif "单位名称" in header_norm:
                safe_write(ws, target_row, col, work.get("单位名称", ""))
            elif "岗位" in header_norm:
                safe_write(ws, target_row, col, work.get("岗位", ""))
            elif "是否邮储银行自主研发工作经验" in header_norm:
                safe_write(ws, target_row, col, work.get("是否邮储银行自主研发工作经验", ""))

def fill_project_experience(ws, project_list):
    keyword_row = find_row_by_keyword(ws, "项目经历（与上述工作经历匹配，仅IT相关经历）")
    if not keyword_row:
        return
    header_row = keyword_row + 1
    data_start = header_row + 1
    reserved = 6
    max_col = ws.max_column

    for r in range(data_start, data_start + reserved):
        clear_row_content(ws, r, max_col)

    sorted_proj = sorted(project_list, key=lambda x: x.get("开始日期", "1900-01-01"), reverse=True)
    need = len(sorted_proj)
    if need > reserved:
        st.warning(f"项目经历共有 {need} 条，但模板只预留了 {reserved} 行，超出部分将被忽略。请手动增加模板预留行数。")
        need = reserved

    for idx in range(need):
        proj = sorted_proj[idx]
        target_row = data_start + idx
        for col in range(1, max_col + 1):
            header = ws.cell(header_row, col).value
            if not header:
                continue
            header_norm = normalize_string(header)
            if "开始日期" in header_norm:
                safe_write(ws, target_row, col, normalize_date(proj.get("开始日期", "")))
            elif "结束日期" in header_norm:
                safe_write(ws, target_row, col, normalize_date(proj.get("结束日期", "")))
            elif "项目名称" in header_norm:
                safe_write(ws, target_row, col, proj.get("项目名称", ""))
            elif "项目描述" in header_norm:
                safe_write(ws, target_row, col, proj.get("项目描述", ""))
            elif "项目角色" in header_norm:
                safe_write(ws, target_row, col, proj.get("项目角色", ""))
            elif "是否邮储银行自主研发工作经验" in header_norm:
                safe_write(ws, target_row, col, proj.get("是否邮储银行自主研发工作经验", ""))

def fill_template(template_path, output_path, ai_data):
    wb = openpyxl.load_workbook(template_path)
    ws = wb.active
    fill_basic_info(ws, ai_data.get("basic", {}))
    if ai_data.get("education", {}).get("undergraduate"):
        fill_education_block(ws, "本科学历", ai_data["education"]["undergraduate"], is_undergrad=True)
    if ai_data.get("education", {}).get("postgraduate"):
        fill_education_block(ws, "研究生学历", ai_data["education"]["postgraduate"], is_undergrad=False)
    if ai_data.get("work_experience"):
        fill_work_experience(ws, ai_data["work_experience"])
    if ai_data.get("project_experience"):
        fill_project_experience(ws, ai_data["project_experience"])
    wb.save(output_path)

# ==================== AI 提取（原有） ====================
SYSTEM_PROMPT = """你是一个专业的简历信息提取助手。请从以下简历文本中提取指定字段，并以 JSON 格式返回。
要求：
1. 日期统一使用 "YYYY-MM-DD" 格式，若只提供年月，则默认补充 "-01"。
2. 若某字段不存在，返回空字符串或空列表。
3. 工作经历和项目经历请按时间由近及远排序（开始日期越晚越靠前）。
4. 最高学历最终输出只有三个选项，本科，硕士，博士，缺省为本科。
5. 是否邮储银行自主研发工作经验输出只有两个选项，是或者否。
6. 掌握语言仅分析计算机相关的开发、测试语言，如java,C语言,Python,Vue等，并非日常交流语言。
7. 专业证书仅分析计算机相关的证书，如果没有则输出无。
8. 对于项目角色的描述不能只是一个简单的岗位名称，还要有相应的工作内容或者责任描述，尽量控制在4条以内。以下文为例：（
    测试工程师：1.负责功能测试、接口测试、自动化测试测试计划与方案制定；
               2.依据系统需求文档，分析业务流程，编写详细的测试计划和方案，明确测试范围、策略、资源安排以及进度计划，确定以等价类划分、边界值分析等方法设计测试用例；
               3.保障功能测试用例设计与执行：针对账户管理、储蓄业务、对公业务、支付结算等核心功能模块；
               4. ...）

输出 JSON 结构如下：
{
    "basic": {
        "姓名": "", "身份证号": "", "出生日期": "", "电话": "",
        "首次参加工作时间": "", "首次参加IT领域工作时间": "", "最高学历": "",
        "掌握语言": "", "掌握技能": "", "专业证书": ""
    },
    "education": {
        "undergraduate": {
            "入学时间": "", "毕业院校": "", "毕业时间": "", "专业": "",
            "毕业证编号": "", "毕业证学信网在线验证码": "",
            "学位证编号": "", "学位证学信网在线验证码": ""
        },
        "postgraduate": {
            "入学时间": "", "毕业院校": "", "毕业时间": "", "专业": "",
            "毕业证编号": "", "毕业证学信网在线验证码": "",
            "学位证编号": "", "学位证学信网在线验证码": ""
        }
    },
    "work_experience": [
        {"开始日期": "", "结束日期": "", "单位名称": "", "岗位": "","是否邮储银行自主研发工作经验": ""}
    ],
    "project_experience": [
        {"开始日期": "", "结束日期": "", "项目名称": "", "项目描述": "", "项目角色": "","是否邮储银行自主研发工作经验": ""}
    ],
    "extra": {"供应商缩写": "", "类型": "", "岗位": "", "级别": ""}
}
"""

def extract_resume_info(api_key: str, full_text: str, model: str = "glm-4-plus") -> dict:
    client = ZhipuAI(api_key=api_key)
    resp = client.chat.completions.create(
        model=model,
        messages=[{"role": "system", "content": SYSTEM_PROMPT},
                  {"role": "user", "content": f"简历文本：\n{full_text}"}],
        temperature=0.1,
        response_format={"type": "json_object"}
    )
    return json.loads(resp.choices[0].message.content)

# ==================== AI 评分与风险分析函数 ====================
SCORE_PROMPT = """你是一位资深 HR 专家，请根据以下简历信息（JSON 格式）和客户岗位技术要求，对该候选人进行综合评分。

评分维度及权重（总分100分）：
1. 学历与专业匹配度（10分）：本科以上学历且专业为计算机相关得高分，如果是工学，理学，管理学和金融方向的专业中等评分，其他学科低分。
2. IT工作年限（10分）：评分跟评定级别挂钩，初级要求是1-3年，中级3-5年，高级5-8年,年限超过评定级别的年限要求越多则该项评分越高（例如中级3年为6分，4年为8分，5年为9分，5年以上为10分满分，不满3年为0分）；如果未提供，则默认8年以上满分，5-8年8分，3-5年5分，3年以下3分。
3. 跳槽频率（10分）：平均每份工作超过18个月满分，12-18个月8分，6-12个月5分，低于6个月2分。
4. 技术栈匹配度（25分）：根据技能或者项目中使用到的技术与岗位要求的关键词重合度评分，同时要注意是否只罗列名词？有无项目中的实际应用描述？谨防“精通泛滥”。如果具有全栈或跨领域能力（如后端+前端、数据库+运维）也是加分项。
**示例分析**：
- 良好：`精通 Spring Boot、MyBatis，熟练使用 MySQL、Redis，了解 Docker、K8s` —— 有明确技术分层。
- 一般：`熟悉 Java、Python、HTML、CSS、JavaScript` —— 过于宽泛，看不出主攻方向。
- 风险：`掌握 SSH 框架` —— 技术过于老旧，可能不满足现代项目要求。
5. 项目经验清晰度（20分）：项目描述是否详细、有量化成果、技术深度。能说清自己负责的模块、使用技术、解决的具体问题，而不能只写“参与xx系统开发”，没有体现个人贡献。以及项目与当前岗位的匹配度，比如对方做的是“电商系统”，和“金融项目”还有差别，则要看技术栈的通用性（如 Spring Cloud、微服务、高并发处理）
6. 项目角色及成果（15分）：是否承担核心角色，是否有具体贡献描述等。是否是大型/复杂项目？规模、并发量、团队构成都是加分项。
7. 稳定性与抗压性（10分）：简历中是否有紧急上线、加班、高强度工作等描述，以及是否有短期工作经历。是否在外包公司工作过？如果已有外包经历，通常能更快适应驻场、需求变更、客户沟通等场景。

请输出 JSON 格式，字段如下：
{
    "scores": {
        "education": 0,
        "work_years": 0,
        "job_hop": 0,
        "tech_match": 0,
        "project_clarity": 0,
        "project_role": 0,
        "stability": 0
    },
    "total": 0,
    "suggestion": "强烈推荐/推荐/待定/不推荐",
    "reason": "简要说明得分依据和推荐理由"
}

简历信息：
{resume_json}

客户岗位技术要求：
{tech_requirements}

评定级别：
{level}
"""

RISK_PROMPT = """你是一位专业的背景调查专家，请根据以下简历信息（JSON 格式），分析该候选人可能存在的潜在风险。

请从以下维度分析（如有则列出，没有则忽略）：
- 职业稳定性风险（频繁跳槽、长期空白期）
- 法律/合规风险（劳动争议、征信问题、学历造假可能）
- 个人因素风险（大龄未婚未育女性可能影响项目稳定性，请注意避免歧视仅作为客观提示、薪资期望过高等）
- 其他隐藏风险（协商离职、外包项目经验不足、跨行业转换等）

输出 JSON 格式，字段如下：
{
    "risks": [
        {
            "category": "稳定性风险",
            "description": "具体描述",
            "level": "高/中/低"
        }
    ],
    "summary": "总体风险评价"
}

简历信息：
{resume_json}
"""

def call_ai_analysis(api_key: str, model: str, prompt: str, resume_json: str, tech_requirements: str = "", level: str = "") -> dict:
    client = ZhipuAI(api_key=api_key)
    if tech_requirements:
        prompt = prompt.replace("{tech_requirements}", tech_requirements)
    else:
        prompt = prompt.replace("{tech_requirements}", "未提供")
    if level:
        prompt = prompt.replace("{level}", level)
    else:
        prompt = prompt.replace("{level}", "未提供")
    prompt = prompt.replace("{resume_json}", resume_json)
    resp = client.chat.completions.create(
        model=model,
        messages=[{"role": "user", "content": prompt}],
        temperature=0.2,
        response_format={"type": "json_object"}
    )
    return json.loads(resp.choices[0].message.content)

# ==================== 新增：基于 RapidOCR 的文字提取和分类 ====================
@st.cache_resource
def init_ocr():
    """初始化 RapidOCR 并缓存"""
    if RAPIDOCR_AVAILABLE:
        return RapidOCR()
    else:
        return None

def extract_text_by_ocr(img_bytes: bytes) -> str:
    """使用 RapidOCR 从图片中提取所有文字，返回合并后的字符串"""
    if not RAPIDOCR_AVAILABLE:
        return ""
    ocr = init_ocr()
    if ocr is None:
        return ""
    try:
        img = Image.open(io.BytesIO(img_bytes))
        img_np = np.array(img)
        result, elapse = ocr(img_np)
        if not result:
            return ""
        texts = []
        for line in result:
            text = line[1][0]
            confidence = line[1][1]
            if confidence > 0.5:
                texts.append(text)
        return " ".join(texts)
    except Exception as e:
        st.warning(f"OCR 识别失败: {e}")
        return ""

def classify_image_by_text(api_key: str, ocr_text: str, img_filename: str) -> str:
    """
    基于 OCR 提取的文字内容，调用大模型判断证件类型
    使用免费的 glm-4-flash 模型
    """
    if not ocr_text.strip():
        return None
    client = ZhipuAI(api_key=api_key)
    types = [
        "身份证正面照片",
        "身份证反面照片",
        "毕业证照片",
        "学位证照片",
        "学信网学历证书电子备案截图",
        "学信网学位证书电子备案截图"
    ]
    options = "\n".join([f"- {t}" for t in types])
    prompt = f"""请根据以下从图片中OCR识别出的文字内容，判断这张图片属于以下哪种证件类型：
{options}

OCR提取的文字内容：
{ocr_text}

区分要点：
- "身份证正面照片"：包含人像、姓名、性别、民族、出生日期、住址、公民身份号码。
- "身份证反面照片"：包含签发机关、有效期限。
- "毕业证照片"：有“毕业证书”字样、学校名称、专业、毕业时间。
- "学位证照片"：有“学位证书”字样、学位级别、学科名称。
- "学信网学历证书电子备案截图"：有“教育部学历证书电子注册备案表”标题，包含姓名、毕业院校、专业、学历层次等。
- "学信网学位证书电子备案截图"：有“中国高等教育学位在线验证报告”标题，包含姓名、学位授予单位、学科名称等。

只输出证件类型名称，不要输出任何其他内容。如果无法确定，输出“未知”。
"""
    try:
        response = client.chat.completions.create(
            model="glm-4-flash",  # 免费模型
            messages=[{"role": "user", "content": prompt}],
            temperature=0.0,
        )
        result = response.choices[0].message.content.strip()
        if result in types:
            return result
        else:
            return None
    except Exception as e:
        st.warning(f"基于文字的分类失败: {e}")
        return None

# ==================== 保留原有的视觉分类函数（备用，未使用但保留） ====================
def classify_image_type(api_key: str, img_bytes: bytes, img_filename: str) -> str:
    """调用智谱视觉模型判断图片属于哪种证件类型（保留原有逻辑）"""
    client = ZhipuAI(api_key=api_key)
    img_base64 = base64.b64encode(img_bytes).decode('utf-8')
    mime_type = "image/jpeg"
    data_url = f"data:{mime_type};base64,{img_base64}"
    
    types = [
        "身份证正面照片",
        "身份证反面照片",
        "毕业证照片",
        "学位证照片",
        "学信网学历证书电子备案截图",
        "学信网学位证书电子备案截图"
    ]
    options = "\n".join([f"- {t}" for t in types])
    prompt = f"""请识别这张图片属于以下哪种证件类型：
{options}

区分要点：
- "身份证正面照片"：包含签发机关、有效期限。
- "身份证反面照片"：包含人像、姓名、性别、民族、出生日期、住址、公民身份号码。
- "毕业证照片"：有“毕业证书”字样、学校名称、专业、毕业时间。
- "学位证照片"：有“学位证书”字样、学位级别、学科名称。
- "学信网学历证书电子备案截图"：有“教育部学历证书电子注册备案表”标题。
- "学信网学位证书电子备案截图"：有“中国高等教育学位在线验证报告”标题。

只输出证件类型名称，不要输出其他内容。如果无法确定，输出“未知”。
"""
    messages = [
        {
            "role": "user",
            "content": [
                {"type": "text", "text": prompt},
                {"type": "image_url", "image_url": {"url": data_url}}
            ]
        }
    ]
    try:
        response = client.chat.completions.create(
            model="glm-4v-plus",
            messages=messages,
            temperature=0.0,
        )
        result = response.choices[0].message.content.strip()
        if result in types:
            return result
        else:
            return None
    except Exception as e:
        st.warning(f"图片 {img_filename} 识别失败: {e}")
        return None

def fill_word_with_images(word_template_bytes, image_classification, new_title=None):
    """将分类好的图片填充到 Word 模板中（适配表格结构：标题行 + 图片行）"""
    doc = Document(io.BytesIO(word_template_bytes))
    # 如果需要修改标题，处理第一个段落
    if new_title:
        if doc.paragraphs:
            first_para = doc.paragraphs[0]
            original_font = None
            if first_para.runs:
                original_font = first_para.runs[0].font
            for run in first_para.runs:
                run.clear()
            new_run = first_para.add_run(new_title)
            if original_font:
                new_run.font.name = original_font.name
                new_run.font.size = original_font.size
                new_run.font.bold = original_font.bold
                new_run.font.italic = original_font.italic
                new_run.font.underline = original_font.underline
                if original_font.color and original_font.color.rgb:
                    new_run.font.color.rgb = original_font.color.rgb
    for title, (img_bytes, _) in image_classification.items():
        found = False
        for table in doc.tables:
            for row_idx, row in enumerate(table.rows):
                for col_idx, cell in enumerate(row.cells):
                    if title in cell.text:
                        if row_idx + 1 < len(table.rows):
                            target_cell = table.cell(row_idx + 1, col_idx)
                        else:
                            target_cell = cell
                        for para in target_cell.paragraphs:
                            drawings = para._element.xpath('.//w:drawing')
                            for draw in drawings:
                                draw.getparent().remove(draw)
                        if target_cell.paragraphs:
                            para = target_cell.paragraphs[0]
                        else:
                            para = target_cell.add_paragraph()
                        run = para.add_run()
                        img_stream = io.BytesIO(img_bytes)
                        run.add_picture(img_stream, width=Inches(5.0))
                        found = True
                        break
                if found:
                    break
            if found:
                break
        if not found:
            st.warning(f"未在表格中找到标题: {title}，请检查模板中的文字是否完全匹配")
    output = io.BytesIO()
    doc.save(output)
    output.seek(0)
    return output.getvalue()

# ==================== Streamlit UI ====================
st.title("📄 智能简历填充工具")
st.markdown("上传 PDF 简历、个人资料图片（身份证、毕业证、学位证、学信网截图等）和 Excel/Word 模板，AI 自动提取信息并生成标准格式文件。")

with st.sidebar:
    st.header("⚙️ 配置")
    api_key = st.text_input("🔑 智谱 AI API Key", type="password")
    model_name = st.selectbox("🤖 选择大模型", ["glm-4-plus", "glm-4-flash", "glm-4-air", "glm-4-long"], index=0)
    st.markdown("---")
    st.caption("Excel 模板需包含：本科学历、研究生学历、工作经历（...）、项目经历（...）等关键字")
    st.caption("Word 模板需包含：身份证正面照片、身份证反面照片、毕业证照片、学位证照片、学信网学历证书电子备案截图、学信网学位证书电子备案截图 等标题")

with st.expander("📌 使用说明", expanded=True):
    st.markdown("""
    1. 准备 **PDF 格式** 的原始简历文件
    2. 准备 **个人资料图片**：身份证正反面、毕业证照片、学位证照片、学信网学历/学位备案表截图等（可选，但强烈推荐）
    3. 准备 **Excel 模板**（.xlsx）或 **Word 模板**（.docx），根据需求选择
    4. 输入你的 **智谱 AI API Key**
    5. 可选填写 **供应商缩写、类型、岗位、级别**（若AI未提取到则以此处为准）
    6. 点击对应按钮进行处理，下载生成的简历文件
    """)

# ---------- 文件上传区域（三列）----------
col1, col2, col3 = st.columns(3)
with col1:
    pdf_file = st.file_uploader("📂 PDF 简历", type=["pdf"])
with col2:
    excel_template = st.file_uploader("📁 Excel 模板", type=["xlsx"])
with col3:
    word_template = st.file_uploader("📄 Word 模板（用于证件照填充）", type=["docx"])

# ---------- 个人资料图片上传区域 ----------
st.markdown("---")
st.subheader("🖼️ 个人资料图片（可选，用于补充学历、身份等信息）")
st.caption("支持上传身份证正反面、毕业证照片、学位证照片、学信网学历/学位备案表截图等，AI 将自动识别图片中的文字并合并到简历中。")
image_files = st.file_uploader(
    "选择图片（可多选）", 
    type=["png", "jpg", "jpeg"], 
    accept_multiple_files=True,
    key="image_uploader"
)

def get_images_hash(images):
    if not images:
        return None
    return hash(tuple((img.name, img.size) for img in images))

current_hash = get_images_hash(image_files)
if current_hash != st.session_state.last_images_hash:
    st.session_state.image_text = None
    st.session_state.uploaded_images = image_files
    st.session_state.last_images_hash = current_hash

if image_files:
    st.write(f"已上传 {len(image_files)} 张图片")
    for img in image_files:
        st.caption(f" - {img.name}")
    
    if st.button("🔍 从图片中提取文字（使用AI视觉模型）", key="ocr_btn"):
        if not api_key:
            st.error("请先在左侧输入智谱AI API Key")
        else:
            with st.spinner(f"正在识别 {len(image_files)} 张图片中的文字，可能需要较长时间..."):
                try:
                    extracted = extract_text_from_images(api_key, model_name, image_files)
                    if extracted.strip():
                        st.session_state.image_text = extracted
                        st.success(f"✅ 成功提取文字，共 {len(extracted)} 字符。")
                        with st.expander("查看提取的文字摘要"):
                            st.text(extracted[:1000] + ("..." if len(extracted) > 1000 else ""))
                    else:
                        st.warning("未从图片中提取到任何文字。")
                        st.session_state.image_text = None
                except Exception as e:
                    st.error(f"❌ 图片识别失败: {e}")

# 检测文件变化，清空 AI 结果
if pdf_file is not None and st.session_state.last_pdf_name != pdf_file.name:
    st.session_state.ai_result = None
    st.session_state.pdf_text = None
    st.session_state.last_pdf_name = pdf_file.name
    st.session_state.score_result = None
    st.session_state.risk_result = None
    st.session_state.level = None
if excel_template is not None and st.session_state.last_excel_name != excel_template.name:
    st.session_state.ai_result = None
    st.session_state.last_excel_name = excel_template.name
    st.session_state.score_result = None
    st.session_state.risk_result = None
    st.session_state.level = None
if word_template is not None and st.session_state.last_word_name != word_template.name:
    st.session_state.word_template = word_template
    st.session_state.last_word_name = word_template.name
    # 不清空 AI 结果，因为 Word 功能独立

# 补充信息输入
with st.container():
    st.subheader("补充信息，用于生成excel命名及人员定级评分")
    col_a, col_b, col_c, col_d = st.columns(4)
    with col_a:
        supplier = st.text_input("供应商缩写 (前4字)", placeholder="例如：北京南天/云南南天")
    with col_b:
        emp_type = st.selectbox("类型", ["研发", "测试"])
    with col_c:
        position = st.selectbox("岗位", ["java开发", "前端开发", "Hadoop", "质量管理", "产品分析", "系统管理", "测试实施", "功能测试", "技术测试"])
    with col_d:
        level = st.selectbox("级别", ["初级", "中级", "高级", "专家"])

# ----------------- Excel 填充处理按钮 -----------------
if st.button("🚀 开始处理 Excel 简历填充", type="primary"):
    missing = []
    if not pdf_file:
        missing.append("PDF 简历")
    if not excel_template:
        missing.append("Excel 模板")
    if not api_key:
        missing.append("智谱 AI API Key")
    if missing:
        st.error(f"❌ 缺少以下必填项：{', '.join(missing)}，请补充后重试。")
        st.stop()

    if st.session_state.pdf_text is None:
        with st.spinner("读取 PDF..."):
            st.session_state.pdf_text = extract_text_from_pdf(pdf_file)
            if not st.session_state.pdf_text.strip():
                st.error("❌ PDF 文本为空，请检查文件是否可解析（例如非扫描件）。")
                st.stop()
    
    full_text = st.session_state.pdf_text
    if st.session_state.image_text:
        full_text += "\n\n【以下为证件/图片中提取的补充信息】\n" + st.session_state.image_text
        st.info(f"已将图片中提取的文字（{len(st.session_state.image_text)} 字符）合并到简历中。")
    else:
        if image_files and st.session_state.image_text is None:
            st.warning("您上传了图片但尚未提取文字，请先点击【从图片中提取文字】按钮，或继续仅使用PDF内容。")
    
    with st.spinner(f"调用 {model_name} 分析简历（可能需要30秒）..."):
        try:
            ai_result = extract_resume_info(api_key, full_text, model_name)
            st.session_state.ai_result = ai_result
            st.success("✅ AI 提取完成")
        except Exception as e:
            st.error(f"❌ AI 调用失败: {e}")
            st.stop()

    extra_hints = []
    if not supplier:
        extra_hints.append("供应商缩写")
    if not emp_type:
        extra_hints.append("类型")
    if not position:
        extra_hints.append("岗位")
    if not level:
        extra_hints.append("级别")
    if extra_hints:
        st.info(f"ℹ️ 以下补充信息未填写，将优先使用 AI 提取结果：{', '.join(extra_hints)}")
    if level != st.session_state.level:
        st.session_state.level = level

# 显示评分与风险分析区域
if st.session_state.ai_result is not None:
    st.markdown("---")
    st.markdown("## 📊 简历智能评估（AI 评分 + 风险分析）")
    
    tech_req = st.text_area(
        "💼 客户岗位技术要求（选输，用于技术匹配度评分）",
        value=st.session_state.tech_requirements,
        placeholder="例如：Java, Spring Boot, MySQL, Redis, 微服务",
        help="输入关键词后点击下方按钮重新评分"
    )
    if tech_req != st.session_state.tech_requirements:
        st.session_state.tech_requirements = tech_req
    
    if st.button("🔍 开始AI评分与风险分析", key="eval_btn"):
        with st.spinner("AI 正在评分及分析风险，请稍候..."):
            try:
                resume_json = json.dumps(st.session_state.ai_result, ensure_ascii=False, indent=2)
                score_result = call_ai_analysis(
                    api_key, model_name, SCORE_PROMPT,
                    resume_json, st.session_state.tech_requirements, st.session_state.level
                )
                st.session_state.score_result = score_result
                risk_result = call_ai_analysis(
                    api_key, model_name, RISK_PROMPT,
                    resume_json, "", ""
                )
                st.session_state.risk_result = risk_result
                st.success("✅ 评估完成")
            except Exception as e:
                st.error(f"❌ AI 评估失败: {e}")
    
    if st.session_state.score_result:
        score = st.session_state.score_result
        st.subheader("📈 评分详情")
        col1, col2, col3 = st.columns(3)
        with col1:
            st.metric("学历专业", f"{score['scores']['education']}/10")
            st.metric("跳槽频率", f"{score['scores']['job_hop']}/10")
            st.metric("项目清晰度", f"{score['scores']['project_clarity']}/20")
        with col2:
            st.metric("IT工作年限", f"{score['scores']['work_years']}/10")
            st.metric("技术栈匹配", f"{score['scores']['tech_match']}/25")
            st.metric("项目角色成果", f"{score['scores']['project_role']}/15")
        with col3:
            st.metric("稳定性/抗压", f"{score['scores']['stability']}/10")
            st.metric("总分", f"{score['total']}/100", delta=None)
        st.info(f"**推荐建议**：{score['suggestion']}  \n**评分理由**：{score['reason']}")
    
    if st.session_state.risk_result:
        st.subheader("⚠️ 潜在风险提示")
        risks = st.session_state.risk_result.get("risks", [])
        if risks:
            for risk in risks:
                level_color = {"高": "🔴", "中": "🟠", "低": "🟡"}.get(risk.get("level", "低"), "⚪")
                st.markdown(f"""<div class="risk-card">
                <strong>{level_color} {risk['category']}</strong>（{risk.get('level', '中')}风险）<br>
                📝 {risk['description']}
                </div>""", unsafe_allow_html=True)
            st.caption(f"📌 总体评价：{st.session_state.risk_result.get('summary', '')}")
        else:
            st.success("✅ 未发现明显风险点")
    
    with st.expander("✏️ 手动补充风险标注（供HR参考）"):
        manual_risk = st.text_area("可输入额外风险备注，如：协商离职、征信问题等", placeholder="例如：候选人上一家公司协商一致离职，需核实背景")
        if manual_risk:
            st.info(f"📝 已记录额外风险：{manual_risk}")

# 显示 AI 提取结果
if st.session_state.ai_result is not None:
    with st.expander("查看 AI 提取结果（结构化数据）", expanded=False):
        st.json(st.session_state.ai_result)

# 下载 Excel 按钮
if st.session_state.ai_result is not None and excel_template is not None:
    extra = st.session_state.ai_result.get("extra", {})
    if supplier:
        extra["供应商缩写"] = supplier
    if emp_type:
        extra["类型"] = emp_type
    if position:
        extra["岗位"] = position
    if level:
        extra["级别"] = level
    st.session_state.ai_result["extra"] = extra

    sup = extra.get("供应商缩写", "未知")[:4]
    name = st.session_state.ai_result.get("basic", {}).get("姓名", "未知")
    typ = extra.get("类型", "研发")
    pos = extra.get("岗位", "java开发")
    lvl = extra.get("级别", "")
    filename = f"{sup}-{name}-{typ}-{pos}-{lvl}-简历.xlsx"
    filename = "".join(c for c in filename if c not in r'\/:*?"<>|')

    if st.button("📥 下载生成的 Excel 简历", type="secondary"):
        with st.spinner("填充 Excel（完全保留原格式）..."):
            with tempfile.NamedTemporaryFile(delete=False, suffix=".xlsx") as tmp:
                tmp.write(excel_template.getbuffer())
                tmp_path = tmp.name
            out_path = os.path.join(tempfile.gettempdir(), filename)
            try:
                fill_template(tmp_path, out_path, st.session_state.ai_result)
                st.success("✅ 填充成功！")
            except Exception as e:
                st.error(f"❌ Excel 填充失败: {e}")
                st.stop()
            finally:
                os.unlink(tmp_path)

        with open(out_path, "rb") as f:
            st.download_button("📥 点击下载文件", f, file_name=filename, mime="application/vnd.openxmlformats-officedocument.spreadsheetml.sheet")
        os.unlink(out_path)

# ----------------- Word 模板图片填充功能 -----------------
if word_template is not None and image_files:
    st.markdown("---")
    st.subheader("📄 Word 证件照自动填充（基于 RapidOCR 文字识别，免费高效）")
    st.caption("系统将使用 RapidOCR 提取图片文字，再根据文字内容自动分类，并填充到 Word 模板对应标题下方。")
    
    if st.button("✨ 开始填充 Word 模板（OCR分类）", key="fill_word_btn_ocr"):
        if not api_key:
            st.error("请先在左侧边栏输入智谱 AI API Key")
        elif not RAPIDOCR_AVAILABLE:
            st.error("RapidOCR 未安装，请运行：pip install rapidocr-onnxruntime")
        else:
            # 从 session_state 中获取必要的命名信息
            if st.session_state.ai_result is not None:
                extra = st.session_state.ai_result.get("extra", {})
                basic = st.session_state.ai_result.get("basic", {})
                sup = extra.get("供应商缩写", supplier)[:4] if not supplier else supplier[:4]
                name = basic.get("姓名", "未知")
                pos = extra.get("岗位", position) if not position else position
                lvl = extra.get("级别", level) if not level else level
            else:
                sup = supplier[:4] if supplier else "未知"
                name = "未知"
                pos = position if position else "java开发"
                lvl = level if level else ""
            
            classified = {}
            progress_bar = st.progress(0)
            total = len(image_files)
            for i, img_file in enumerate(image_files):
                img_bytes = img_file.getvalue()
                ocr_text = extract_text_by_ocr(img_bytes)
                if not ocr_text.strip():
                    st.warning(f"图片 {img_file.name} OCR 未提取到文字，跳过")
                else:
                    img_type = classify_image_by_text(api_key, ocr_text, img_file.name)
                    st.write(f"图片 {img_file.name} -> OCR 文字片段：{ocr_text[:100]}... -> 分类为: {img_type}")
                    if img_type:
                        classified[img_type] = (img_bytes, img_file.name)
                    else:
                        st.warning(f"无法根据文字识别图片 {img_file.name} 的证件类型，已跳过")
                progress_bar.progress((i+1)/total)
            progress_bar.empty()
            
            if not classified:
                st.error("未能识别任何有效证件图片，请确保图片清晰且包含足够的文字信息。")
            else:
                with st.spinner("正在填充 Word 模板并修改标题..."):
                    try:
                        word_bytes = word_template.getvalue()
                        new_title = f"{sup}-{name}-{pos}-{lvl}-个人资料"
                        output_bytes = fill_word_with_images(word_bytes, classified, new_title=new_title)
                        word_filename = f"{sup}-{name}-{pos}-{lvl}-资料.docx"
                        word_filename = "".join(c for c in word_filename if c not in r'\/:*?"<>|')
                        st.success("✅ Word 填充成功！")
                        st.download_button(
                            label="📥 下载填充后的 Word 文件",
                            data=output_bytes,
                            file_name=word_filename,
                            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                        )
                    except Exception as e:
                        st.error(f"❌ Word 处理失败: {e}")
else:
    if word_template is None and image_files:
        st.info("💡 若需使用 Word 填充功能，请先上传 Word 模板。")
    elif word_template and not image_files:
        st.info("💡 若需使用 Word 填充功能，请先在上方“个人资料图片”区域上传证件照片。")
